import gradio as gr
from httpx import AsyncClient
from docx import Document
import numpy as np
# 在WebUI类初始化前配置HTTP客户端
gr.routes.client = AsyncClient(verify=False)
import time
from pdf_processor import PDFProcessor
from vector_db import VectorDB
from llm_integration import PatentGenerator
import os
os.environ["SSL_CERT_FILE"] = r"H:\anadonda\envs\OpenPatent\Library\ssl\cacert.pem"
from docx.oxml.ns import qn

class WebUI:
    """
    网页用户界面类，用于创建和管理专利生成系统的用户界面。
    """
    def __init__(self):
        """
        初始化网页用户界面。
        """
        self.patent_generator = PatentGenerator()
        self.current_stage = None
        self.current_doc_type = None
        self.db_paths = {
            '摘要': r'dbs\abstract',
            '说 明 书': r'dbs\specification',
            '权 利 要 求 书': r'dbs\claims'
        }
        self.use_existing_db = False

    def init_interface(self):
        """
        初始化用户界面。

        返回:
        gr.Blocks: Gradio 界面块。
        """
        with gr.Blocks(title="OpenPatent 专利生成系统") as demo:
            gr.Markdown("## OpenPatent 专利文档生成系统")
            
            with gr.Tab("1. 选择参考专利"):
                ref_patents = gr.Files(label="上传参考专利文件(PDF)")
                process_btn = gr.Button("处理专利文件")
                process_btn2 = gr.Button("已有本地知识库，点击这里")
                process_output = gr.Markdown()
            
            with gr.Tab("2. 上传技术文档"):
                tech_doc = gr.File(label="技术文档(.docx)")
                
            with gr.Tab("3. 生成专利文档"):
                with gr.Row():
                    gen_spec_btn = gr.Button("生成说明书")
                    gen_abstract_btn = gr.Button("生成摘要")
                    gen_claims_btn = gr.Button("生成权利要求书")
                
                with gr.Column():
                    output_preview = gr.Chatbot(label="专利生成过程", height=500, elem_id="centered-chat")
                
                with gr.Row():
                    user_feedback = gr.Textbox(label="修改意见", lines=3)
                    submit_feedback = gr.Button("提交反馈", variant="primary")
            
            # 绑定事件
            process_btn.click(self.process_patents, inputs=ref_patents, outputs=process_output)
            process_btn2.click(self.load_existing_db, outputs=process_output)
            gen_spec_btn.click(self.generate_specification, inputs=tech_doc, outputs=output_preview)
            gen_abstract_btn.click(self.generate_abstract, inputs=tech_doc, outputs=output_preview)
            gen_claims_btn.click(self.generate_claims, inputs=tech_doc, outputs=output_preview)
            submit_feedback.click(self.submit_feedback, inputs=user_feedback, outputs=output_preview)
            
        return demo

    def process_patents(self, files):
        """
        处理上传的参考专利文件，创建向量数据库并保存索引。

        参数:
        files (list): 上传的参考专利文件列表。

        返回:
        str: 处理结果信息。
        """
        if not files:
            return "请上传参考专利文件"
        processor = PDFProcessor()
        
        section_array = []
        for file in files:
            sections = processor.split_pdf(file.name)
            for db_type in self.db_paths:
                section_content = sections.get(db_type)
                section_array.append(section_content)
                if section_content:
                    print(f"dbtype:{db_type}")
                    print(section_content)
                else:
                    print(f"Section {db_type} not found in PDF")
        section_array = np.array(section_array)
        section_array = section_array.reshape((-1,3))
        print(f"翻转前：{section_array}")
        section_array = section_array.T
        print(f"翻转后：{section_array}")
        self.db_list = [0 for i in range(3)]
        for i,db_type in enumerate(self.db_paths):
            self.db_list[i] = VectorDB(db_type)
            self.db_list[i].create_index(section_array[i])
            self.db_list[i].save_index(self.db_paths[db_type])
        self.use_existing_db = True
        return "参考专利处理完成，已建立三个知识库！"

    def load_existing_db(self):
        """
        加载已有的本地知识库。

        返回:
        str: 加载结果信息。
        """
        self.use_existing_db = True
        self.db_list = [0 for i in range(3)]
        for i, db_type in enumerate(self.db_paths):
            self.db_list[i] = VectorDB(db_type)
            self.db_list[i].load_index(self.db_paths[db_type])
        return "已加载本地知识库"

    def generate_specification(self, tech_doc):
        """
        生成专利说明书。

        参数:
        tech_doc: 上传的技术文档。

        返回:
        list: 包含系统消息和生成内容的列表。
        """
        self.current_stage = "specification"
        return self._generate_draft(tech_doc, "说 明 书", "说明书")

    def generate_abstract(self, tech_doc):
        """
        生成专利摘要。

        参数:
        tech_doc: 上传的技术文档。

        返回:
        list: 包含系统消息和生成内容的列表。
        """
        return self._generate_draft(tech_doc, "摘要", "摘要")

    def generate_claims(self, tech_doc):
        """
        生成专利权利要求书。

        参数:
        tech_doc: 上传的技术文档。

        返回:
        list: 包含系统消息和生成内容的列表。
        """
        return self._generate_draft(tech_doc, "权 利 要 求 书", "权利要求书")

    def _generate_draft(self, tech_doc, db_type: str, doc_type: str):
        """
        生成专利文档初稿。

        参数:
        tech_doc: 上传的技术文档。
        db_type (str): 数据库类型，如 "摘要", "说明书", "权利要求书"。
        doc_type (str): 文档类型，如 "说明书", "摘要", "权利要求书"。

        返回:
        list: 包含系统消息和生成内容的列表。
        """
        if not self.use_existing_db:
            return [("系统", "请先处理参考专利或选择已有知识库")]
        if tech_doc is None:
            return [("系统", "请先上传技术文档")]

        # 读取技术文档内容
        doc = Document(tech_doc.name)
        query = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

        # 加载向量数据库
        db_map = {"摘要":0,"说 明 书":1,"权 利 要 求 书":2}
        vector_db = self.db_list[db_map[db_type]]
        # 检索相关内容
        related_patents = vector_db.query(query, top_k=2)
        print(related_patents)
        context = "\n".join(related_patents) if related_patents else "无相关专利内容"
        
        # 生成专利文档
        try:
            content = self.patent_generator.generate_draft(query, context, doc_type)
            self.current_doc_type = doc_type
            return [
                ("系统", "开始生成专利文档..."),
                ("助手", content)
            ]
        except Exception as e:
            return [
                ("系统", "开始生成专利文档..."),
                ("助手", f"生成失败: {str(e)}")
            ]

    def submit_feedback(self, feedback):
        """
        提交用户反馈，根据反馈内容进行文档保存或修订。

        参数:
        feedback (str): 用户的反馈意见。

        返回:
        list: 包含系统消息和处理结果的列表。
        """
        if not self.current_doc_type:
            return [("系统", "请先生成草案")]
        if not feedback.strip():
            return [("系统", "请输入修改意见")]

        try:
            if '满意' in feedback:
                current_content = self.patent_generator.current_draft.get(self.current_doc_type, "No draft available")
                messages = [
                    ("系统", "文档已确认满意，开始保存..."),
                    ("助手", current_content)
                ]
                from docx import Document
                from docx.shared import Pt
                from docx.enum.text import WD_LINE_SPACING

                # 创建新文档并设置全局样式
                doc = Document()
                doc.styles['Normal'].font.name = '宋体'
                doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                # 确保 Heading 1 样式的字体为宋体
                heading_style = doc.styles['Heading 1']
                heading_style.font.name = '宋体'
                heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                heading_style.font.size = Pt(10.5)  # 宋体五号
                heading_style.font.bold = True
                heading_style.paragraph_format.space_before = Pt(6)
                heading_style.paragraph_format.space_after = Pt(6)
                heading_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

                # 创建一个新的段落样式用于正文
                body_style = doc.styles['Normal']
                body_style.font.name = '宋体'
                body_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                body_style.font.size = Pt(10.5)  # 宋体五号
                body_style.paragraph_format.space_before = Pt(0)  # 正文段落前无缩进
                body_style.paragraph_format.space_after = Pt(0)   # 正文段落后无缩进
                body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                body_style.paragraph_format.first_line_indent = Pt(0)  # 无缩进
                # 解析XML标签并应用格式
                paragraphs = current_content.split('\n')
                for para in paragraphs:
                    if para.startswith('<标题>'):
                        title = para[4:-5].strip()  # 提取<标题>内容</标题>
                        title_para = doc.add_paragraph(title, style='Heading 1')
                    elif para.startswith('<段落>'):
                        content = para[4:-5].strip()  # 提取<段落>内容</段落>
                        p = doc.add_paragraph(content)
                        # p.paragraph_format.space_before = Pt(3)
                        # p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        # p.paragraph_format.first_line_indent = Pt(0)  # 无缩进

                # 保存文档
                filename = f"{self.current_doc_type}.gradio_{int(time.time())}.docx"
                doc.save(filename)
                messages.append( ("系统", f"文件已保存为：{filename}") )
                return messages
            else:
                revised_content = self.patent_generator.revise_draft(feedback, self.current_doc_type)
                messages = [
                    ("系统", "开始修订文档..."),
                    ("助手", revised_content)
                ]
                return messages
        except Exception as e:
            return [
                ("系统", "处理反馈时发生错误"),
                ("助手", f"错误详情: {str(e)}")
            ]

if __name__ == "__main__":
    WebUI().init_interface().launch()
